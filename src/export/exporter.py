"""
Export transcript and notes to various formats
"""
from pathlib import Path
from datetime import datetime


def _format_time(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def _build_transcript_text(segments: list, include_timestamps=True) -> str:
    lines = []
    current_speaker = None
    for seg in segments:
        speaker = seg.speaker or "講者"
        if speaker != current_speaker:
            lines.append(f"\n{speaker}:")
            current_speaker = speaker
        ts = f"[{_format_time(seg.start)}] " if include_timestamps else ""
        lines.append(f"  {ts}{seg.text}")
    return "\n".join(lines).strip()


def export_txt(segments: list, notes: str, title: str, output_path: str):
    """Export as plain text."""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"# {title}\n")
        f.write(f"匯出時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 60 + "\n\n")

        f.write("## 逐字稿\n\n")
        f.write(_build_transcript_text(segments))
        f.write("\n\n")

        if notes:
            f.write("=" * 60 + "\n")
            f.write("## 整理筆記\n\n")
            f.write(notes)


def export_docx(segments: list, notes: str, title: str, output_path: str):
    """Export as Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"匯出時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # Transcript section
    doc.add_heading("逐字稿", level=1)

    current_speaker = None
    for seg in segments:
        speaker = seg.speaker or "講者"
        if speaker != current_speaker:
            p = doc.add_paragraph()
            run = p.add_run(f"{speaker}：")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
            current_speaker = speaker

        p = doc.add_paragraph(style='List Bullet')
        ts_run = p.add_run(f"[{_format_time(seg.start)}] ")
        ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        ts_run.font.size = Pt(9)
        p.add_run(seg.text)

    # Notes section
    if notes:
        doc.add_page_break()
        doc.add_heading("整理筆記", level=1)

        # Simple markdown-to-docx conversion
        for line in notes.split('\n'):
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                # Handle **bold**
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.bold = (i % 2 == 1)

    doc.save(output_path)


def export_pdf(segments: list, notes: str, title: str, output_path: str):
    """Export as PDF."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()

    # Use built-in font (supports basic Latin; for CJK use a TTF font)
    # Try to load a CJK font if available
    try:
        import platform
        if platform.system() == 'Darwin':
            font_path = "/System/Library/Fonts/STHeiti Light.ttc"
        else:
            font_path = "C:/Windows/Fonts/msyh.ttc"

        pdf.add_font('CJK', '', font_path, uni=True)
        pdf.set_font('CJK', size=16)
    except Exception:
        pdf.set_font('Helvetica', size=16)

    # Title
    pdf.cell(0, 10, title, ln=True, align='C')
    pdf.set_font_size(10)
    pdf.cell(0, 8, f"Export: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
    pdf.ln(5)

    # Transcript
    pdf.set_font_size(14)
    pdf.cell(0, 10, "Transcript", ln=True)
    pdf.set_font_size(10)

    current_speaker = None
    for seg in segments:
        speaker = seg.speaker or "Speaker"
        if speaker != current_speaker:
            pdf.set_font_size(11)
            pdf.set_text_color(26, 115, 232)
            pdf.cell(0, 8, speaker, ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font_size(10)
            current_speaker = speaker

        ts = _format_time(seg.start)
        line = f"  [{ts}] {seg.text}"
        pdf.multi_cell(0, 6, line)

    # Notes
    if notes:
        pdf.add_page()
        pdf.set_font_size(14)
        pdf.cell(0, 10, "Notes", ln=True)
        pdf.set_font_size(10)
        # Strip markdown formatting for PDF
        clean_notes = notes.replace('**', '').replace('*', '')
        for line in clean_notes.split('\n'):
            if line.startswith('#'):
                pdf.set_font_size(12)
                pdf.cell(0, 8, line.lstrip('#').strip(), ln=True)
                pdf.set_font_size(10)
            else:
                pdf.multi_cell(0, 6, line if line else ' ')

    pdf.output(output_path)
